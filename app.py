import os
import io
import uuid  # Library for generating unique IDs
from flask import Flask, render_template, request, jsonify, send_file, session
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document

# --- SETUP ---
# Load environment variables from a .env file for security
load_dotenv()
app = Flask(__name__)
# A secret key is required for Flask sessions to work, enabling user-specific data.
app.secret_key = os.urandom(24)

# Initialize the OpenAI client. 
# It automatically uses the OPENAI_API_KEY from your .env file.
try:
    client = OpenAI()
except Exception as e:
    print(f"Error: Could not initialize OpenAI client. Check your API key. Details: {e}")
    client = None

# --- DEFAULT QUESTIONS ---
# This list will be used if no custom questions are uploaded.
DEFAULT_QUESTIONS = [
    "How did you get to know the Subject?",
    "How did Subject conduct the sales process?/How did Subject pitch the product in question to you?",
    "What was shared with you during the sales process?",
    "How many times did you meet the Subject before purchasing the policy? How long was each session?",
    "What are your financial goals communicated to the Subject?",
    "Did you approach the Subject or did the Subject approach you for the policy purchased?",
    "Did you receive the POS documents after the signing of KYC?",
    "How did you come to realise the issue?",
    "Did you raise the issue to the Subject? Did the Subject clarify the issues faced? What was shared during the clarification?",
    "Do you have any correspondence to support your statement?"
]

# --- HELPER FUNCTION TO GET QUESTIONS ---
def get_questions_for_session():
    """ Safely gets the question list from the session, falling back to default. """
    return session.get('questions', DEFAULT_QUESTIONS)

# --- NEW: UPLOAD AND RESET ROUTES ---
@app.route('/upload_questions', methods=['POST'])
def upload_questions():
    """ Handles the upload of a .docx file with custom questions. """
    if 'question_file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files['question_file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    if file and file.filename.endswith('.docx'):
        try:
            document = Document(file.stream)
            # Read each paragraph as a question, ignoring empty ones.
            questions = [p.text for p in document.paragraphs if p.text.strip()]
            if not questions:
                return jsonify({"error": "Word document is empty or contains no text."}), 400
            
            # Store the custom questions in the user's session.
            session['questions'] = questions
            return jsonify({"success": True, "filename": file.filename, "question_count": len(questions)})
        except Exception as e:
            return jsonify({"error": f"Error parsing Word document: {e}"}), 500
    else:
        return jsonify({"error": "Invalid file type. Please upload a .docx file."}), 400

@app.route('/reset', methods=['POST'])
def reset_session():
    """ Clears custom questions from the session, reverting to default. """
    session.pop('questions', None)
    return jsonify({"success": True, "message": "Session reset."})


# --- MAIN ROUTE ---
@app.route('/')
def index():
    """ Renders the main user interface from the index.html template. """
    return render_template('index.html')

# --- REAL-TIME PROCESSING ROUTES (MODIFIED) ---
@app.route('/transcribe', methods=['POST'])
def transcribe_audio():
    """ Handles transcription of audio chunks with a UNIQUE filename for each request. """
    if client is None: return jsonify({"error": "OpenAI client is not initialized."}), 500
    if 'audio_data' not in request.files: return jsonify({"error": "No audio file provided."}), 400
    audio_file = request.files['audio_data']
    unique_filename = f"temp_{uuid.uuid4()}.webm"
    temp_audio_path = unique_filename
    audio_file.save(temp_audio_path)
    try:
        with open(temp_audio_path, "rb") as f:
            transcript = client.audio.transcriptions.create(model="whisper-1", file=f, language="en")
        os.remove(temp_audio_path)
        return jsonify({'transcript': transcript.text})
    except Exception as e:
        if os.path.exists(temp_audio_path): os.remove(temp_audio_path)
        print(f"Error during transcription: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/realtime_feedback', methods=['POST'])
def realtime_feedback():
    """ Analyzes the current transcript using session questions and provides a coaching suggestion. """
    if client is None: return jsonify({"error": "OpenAI client not initialized."}), 500
    transcript = request.json.get('transcript')
    if not transcript:
        return jsonify({'feedback': 'Start by introducing yourself and the purpose of the interview.'})

    questions = get_questions_for_session()
    question_list_str = "\n".join(f"- {q}" for q in questions)
    
    system_prompt = f"""You are an AI assistant coaching an interviewer in real-time. Your goal is to help them conduct a comprehensive interview based on a list of required questions.
**Context:**
- You will receive the [CURRENT TRANSCRIPT] of the conversation so far.
- Your task is to check which questions from the [QUESTION LIST] have NOT been covered.
- Provide a single, concise, and helpful suggestion for the interviewer in under 20 words.
- Your feedback should gently guide the interviewer to the next uncovered topic.
- If the last statement was vague, suggest clarification.
**[QUESTION LIST]:**
{question_list_str}

**Your Task:**
Based on the [CURRENT TRANSCRIPT], what is the best *single* suggestion for the interviewer right now?
"""
    try:
        response = client.chat.completions.create(model="gpt-4o", messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"[CURRENT TRANSCRIPT]:\n{transcript}"}])
        return jsonify({'feedback': response.choices[0].message.content})
    except Exception as e:
        print(f"Error during feedback generation: {e}")
        return jsonify({'feedback': '...'}), 200

# --- FINAL PROCESSING ROUTES (MODIFIED) ---
@app.route('/summarize', methods=['POST'])
def summarize_text():
    """ Takes the final transcript and generates the summary using session questions. """
    if client is None: return jsonify({"error": "OpenAI client is not initialized."}), 500
    transcript = request.json.get('transcript')
    if not transcript: return jsonify({"error": "No transcript text provided"}), 400

    questions = get_questions_for_session()
    question_list_str = "\n".join(f"- {q}" for q in questions)

    system_prompt = f"""You are an AI assistant for a Market Conduct investigation. Your task is to create a structured summary of an interview transcript.
Instructions:
1. Read the entire provided [TRANSCRIPT].
2. Answer each of the questions from the list below based *only* on the information in the transcript.
3. If the answer to a question cannot be found, you MUST explicitly state: "Information regarding this was not discussed in the interview."
4. Quote the complainant directly where necessary to support an answer.
5. Format the output clearly with each question on a new line followed by its answer.
Questions to Answer:
{question_list_str}
"""
    try:
        response = client.chat.completions.create(model="gpt-4o", messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Here is the transcript:\n\n[TRANSCRIPT]\n{transcript}"}])
        return jsonify({'summary': response.choices[0].message.content})
    except Exception as e:
        print(f"Error during summarization: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/download_summary', methods=['POST'])
def download_summary():
    """ Receives the final summary and transcript, and returns them as a Word document. """
    data = request.get_json()
    summary_text = data.get('summary')
    transcript_text = data.get('transcript')
    if not summary_text: return "No summary content provided", 400
    document = Document()
    document.add_heading('Interview Report', level=1)
    document.add_heading('Generated Summary', level=2)
    document.add_paragraph(summary_text)
    document.add_page_break()
    document.add_heading('Full Interview Transcript', level=2)
    if transcript_text: document.add_paragraph(transcript_text)
    else: document.add_paragraph("Transcript was not provided for the download.")
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return send_file(
        file_stream,
        as_attachment=True,
        download_name='Interview_Report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# --- RUN THE APP ---
if __name__ == '__main__':
    app.run(debug=True)
